"""Safe, local text extraction for the supported course-material formats."""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader


MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_DOCX_DOCUMENT_XML_BYTES = 50 * 1024 * 1024
CHUNK_SIZE = 1_200
CHUNK_OVERLAP = 180


class DocumentParseError(ValueError):
    """Raised when an uploaded file cannot safely yield readable text."""


@dataclass(frozen=True)
class ExtractedSection:
    content: str
    page_number: int | None


@dataclass(frozen=True)
class ParsedChunk:
    content: str
    page_number: int | None


def validate_file_signature(path: Path, file_type: str) -> None:
    """Reject common extension spoofing before a parser opens the file."""
    if file_type == "pdf":
        with path.open("rb") as uploaded_file:
            header = uploaded_file.read(1_024)

        if b"%PDF-" not in header:
            raise DocumentParseError("PDF 文件头无效，文件可能不是 PDF。")

    elif file_type == "docx":
        if not zipfile.is_zipfile(path):
            raise DocumentParseError("DOCX 文件格式无效，文件可能已损坏或被伪装。")

        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            uncompressed_size = sum(
                entry.file_size
                for entry in archive.infolist()
            )
            document_info = archive.getinfo("word/document.xml") if "word/document.xml" in names else None

            if document_info is None:
                raise DocumentParseError("DOCX 缺少正文内容。")

            if (
                uncompressed_size > MAX_DOCX_UNCOMPRESSED_BYTES
                or document_info.file_size > MAX_DOCX_DOCUMENT_XML_BYTES
            ):
                raise DocumentParseError("DOCX 解压后的内容过大，已拒绝解析。")


def parse_document(path: Path, file_type: str) -> list[ParsedChunk]:
    """Extract text locally and split it into searchable, source-aware chunks."""
    validate_file_signature(path, file_type)

    if file_type == "pdf":
        sections = _extract_pdf(path)
    elif file_type == "docx":
        sections = _extract_docx(path)
    elif file_type in {"txt", "md"}:
        sections = [_extract_text_file(path)]
    else:
        raise DocumentParseError("不支持的文件类型。")

    chunks: list[ParsedChunk] = []
    for section in sections:
        for content in _split_text(section.content):
            chunks.append(
                ParsedChunk(
                    content=content,
                    page_number=section.page_number,
                )
            )

    if not chunks:
        raise DocumentParseError(
            "无法从文件中提取可检索文字；扫描版 PDF 暂不支持 OCR。"
        )

    return chunks


def _extract_pdf(path: Path) -> list[ExtractedSection]:
    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # pypdf exposes several parser-specific errors.
        raise DocumentParseError("PDF 已损坏或无法读取。") from exc

    sections: list[ExtractedSection] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            content = _normalize_text(page.extract_text() or "")
        except Exception as exc:
            raise DocumentParseError(
                f"PDF 第 {page_number} 页无法提取文字。"
            ) from exc

        if content:
            sections.append(
                ExtractedSection(
                    content=content,
                    page_number=page_number,
                )
            )

    return sections


def _extract_docx(path: Path) -> list[ExtractedSection]:
    try:
        document = Document(str(path))
    except Exception as exc:
        raise DocumentParseError("DOCX 已损坏或无法读取。") from exc

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    content = _normalize_text("\n".join(paragraphs))
    return [ExtractedSection(content=content, page_number=None)] if content else []


def _extract_text_file(path: Path) -> ExtractedSection:
    raw_bytes = path.read_bytes()

    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            content = raw_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise DocumentParseError("TXT/MD 仅支持 UTF-8 或 GB18030 编码。")

    return ExtractedSection(
        content=_normalize_text(content),
        page_number=None,
    )


def _normalize_text(content: str) -> str:
    content = content.replace("\x00", "")
    content = content.replace("\r\n", "\n").replace("\r", "\n")
    content = re.sub(r"[\t\f\v ]+", " ", content)
    content = re.sub(r"\n{3,}", "\n\n", content)
    return content.strip()


def _split_text(content: str) -> list[str]:
    normalized = _normalize_text(content)
    if not normalized:
        return []

    chunks: list[str] = []
    start = 0
    content_length = len(normalized)

    while start < content_length:
        end = min(start + CHUNK_SIZE, content_length)

        if end < content_length:
            boundary = max(
                normalized.rfind(marker, start + CHUNK_SIZE // 2, end)
                for marker in ("\n", "。", "！", "？", ".", ";", "；", " ")
            )
            if boundary > start:
                end = boundary + 1

        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= content_length:
            break

        start = max(end - CHUNK_OVERLAP, start + 1)

    return chunks
